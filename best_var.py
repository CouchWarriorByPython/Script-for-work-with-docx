from docx import Document
from datetime import datetime, timedelta


def check_work_shift():
    enter = input('Укажите смену, дневная/ночна(д/н): ')
    try:
        if enter.lower() == 'дневная' or enter.lower() == 'д':
            data = datetime.today()
            return data.strftime('%d.%m.%Y')
        else:
            data = datetime.today() + timedelta(days=1)
            return data.strftime('%d.%m.%Y')

    except Exception as ex:
        print('Некорректный ввод', ex)


def func(table):
    list_full = []
    list_short = []
    count = 0
    for row in table.rows[2:]:
        count += 1
        string = ''
        for cell in row.cells[1:3]:
            i = cell.text.rstrip('\n ').strip(' ')
            string = f"{string + i + ', '}"
        string = string.rstrip(', ')
        new_line = f"-\t{string[0].lower() + string[1:]}"
        list_full.append(new_line)

    for row in table.rows[2:]:
        string = ''
        for cell in row.cells[1:2]:
            i = cell.text.rstrip('\n ').strip(' ')
            string = f"{i};"
        new_line = f"-\t{string[0].lower() + string[1:]}"
        list_short.append(new_line)

    writer(list_full, list_short, count)


def check():
    while True:
        try:
            doc = Document('1.docx').tables[0]
            func(doc)
            break

        except Exception as ex:
            print('Некорректный ввод, проверьте путь к вашему файлу\n', ex, sep='')


def writer(list_fl, list_st, ct):
    while True:
        try:
            doc = Document('1.docx')
            par = [par._element.getparent().remove(par._element) for par in doc.paragraphs]
            table = [tab._element.getparent().remove(tab._element) for tab in doc.tables]

            print('Введите шапку')
            space = input('> ')

            print('Введите номер предписания:')
            number_ceh = input('> ')

            date_time = check_work_shift()
            doc.add_paragraph(space).runs[0].bold = True
            doc.add_paragraph(f'Выявлено {ct} нарушений\n')
            for par_fl in list_fl:
                par_fl = f'\n{par_fl}\n'
                doc.add_paragraph(par_fl)

            doc.add_paragraph(f'Выдано предписание № {number_ceh}-1 от {date_time}\n').runs[0].bold = True

            for par_st in list_st[:-1]:
                par_st = f'{par_st}'
                doc.add_paragraph(par_st)

            doc.add_paragraph(f'{list_st[-1]}.'.replace(';', ''))
            doc.save('1.docx')
            break
        except Exception as ex:
            print('Некорректный ввод, проверьте название вашего файла\n', ex, sep='')


if __name__ == '__main__':
    check()
